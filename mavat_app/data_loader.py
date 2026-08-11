from __future__ import annotations

import csv
import re
from pathlib import Path
from typing import Any


HEADER_ALIASES = {
    "שם": "client_name",
    "שם לקוח": "client_name",
    "לקוח": "client_name",
    "תז": "id_number",
    "ת.ז": "id_number",
    "תעודת זהות": "id_number",
    "מספר זהות": "id_number",
    "גוש": "block",
    "חלקה": "parcel",
    "מגרש": "lot",
    "ישוב": "locality",
    "יישוב": "locality",
    "מספר תכנית": "plan_number",
    "מספר תוכנית": "plan_number",
    "שם תכנית": "plan_name",
    "שם תוכנית": "plan_name",
    "ועדה": "committee",
    "ועדה מקומית": "committee",
    "מרחב תכנון": "planning_area",
    "סוג תכנית": "plan_type",
    "סוג תוכנית": "plan_type",
    "שטח": "area",
    "שטח בדונם": "area",
}


def _clean_header(value: Any, index: int) -> str:
    text = re.sub(r"\s+", " ", str(value or "").strip())
    text = text.replace('ת"ז', "תז").replace("ת׳ז", "תז")
    return HEADER_ALIASES.get(text, text or f"עמודה_{index + 1}")


def _rows_to_records(rows: list[list[Any]]) -> list[dict[str, Any]]:
    rows = [list(row) for row in rows if any(str(v or "").strip() for v in row)]
    if not rows:
        return []
    headers = [_clean_header(value, i) for i, value in enumerate(rows[0])]
    seen: dict[str, int] = {}
    unique_headers: list[str] = []
    for header in headers:
        seen[header] = seen.get(header, 0) + 1
        unique_headers.append(header if seen[header] == 1 else f"{header}_{seen[header]}")

    records: list[dict[str, Any]] = []
    for row_number, row in enumerate(rows[1:], start=2):
        record = {
            unique_headers[i]: (row[i] if i < len(row) and row[i] is not None else "")
            for i in range(len(unique_headers))
        }
        record["_row_number"] = row_number
        records.append(record)
    return records


def load_xlsx(path: Path) -> list[dict[str, Any]]:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError("יש להתקין openpyxl כדי לקרוא קובצי Excel") from exc
    workbook = load_workbook(path, read_only=True, data_only=True)
    try:
        sheet = workbook.active
        return _rows_to_records([list(row) for row in sheet.iter_rows(values_only=True)])
    finally:
        workbook.close()


def load_csv(path: Path) -> list[dict[str, Any]]:
    raw = path.read_bytes()
    text = None
    for encoding in ("utf-8-sig", "cp1255", "utf-16"):
        try:
            text = raw.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    if text is None:
        raise ValueError("לא ניתן לזהות את קידוד קובץ ה-CSV")
    sample = text[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t")
    except csv.Error:
        dialect = csv.excel
    return _rows_to_records([list(row) for row in csv.reader(text.splitlines(), dialect)])


def load_docx(path: Path) -> list[dict[str, Any]]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("יש להתקין python-docx כדי לקרוא קובצי Word") from exc
    document = Document(path)
    for table in document.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if rows:
            return _rows_to_records(rows)

    lines = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    if not lines:
        return []
    delimiter = "\t" if "\t" in lines[0] else (";" if ";" in lines[0] else ",")
    return _rows_to_records([line.split(delimiter) for line in lines])


def load_records(file_path: str | Path) -> list[dict[str, Any]]:
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(path)
    extension = path.suffix.lower()
    if extension == ".xlsx":
        return load_xlsx(path)
    if extension in {".csv", ".tsv"}:
        return load_csv(path)
    if extension == ".docx":
        return load_docx(path)
    raise ValueError("סוג קובץ לא נתמך. יש לבחור XLSX, CSV, TSV או DOCX")
